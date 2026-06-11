# -*- coding: utf-8 -*-
"""Read interview_data.json and generate a Word document."""
import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = Path(__file__).resolve().parent.parent

doc = Document()

# -- Global style --
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
style.paragraph_format.line_spacing = 1.35
style.paragraph_format.space_after = Pt(4)

# Title
title = doc.add_heading("AI Chat Agent 项目面试问题与答案", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("基于个人开源项目 my_AI_chat 的面试准备文档").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("涵盖：Agent 架构、RAG/GraphRAG、MCP 协议、工程化部署、LLM 集成等").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("")


def add_question(num, text):
    p = doc.add_paragraph()
    run = p.add_run("Q{}：{}".format(num, text))
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0, 70, 130)
    p.paragraph_format.space_before = Pt(18)


def add_answer(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Cm(0.5)


# Load data
data = json.loads((BASE / "data" / "interview_data.json").read_text(encoding="utf-8"))
q_counter = [0]  # mutable counter

for section in data["sections"]:
    doc.add_heading(section["title"], level=1)
    for item in section["questions"]:
        q_counter[0] += 1
        add_question(q_counter[0], item["q"])
        add_answer(item["a"])

# Save
output = BASE / "exports" / "项目面试问题.docx"
doc.save(str(output))
print("Document saved to: {}".format(output))
